from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponseRedirect
from .forms import *
from .models import *
import pandas as pd
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
from django.db.models import Count
import json
from django.contrib import messages
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.db.models import Count, Q
import pdfkit
from docx import Document
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.decorators import login_required, permission_required, user_passes_test
from django.contrib.auth import logout
from django.db.models import Case, When, Value, BooleanField, Exists, OuterRef


def is_gestionnaire(user):
    return user.groups.filter(name='gestionnaire').exists()

def is_admin(user):
    return user.is_superuser or user.groups.filter(name='admin').exists()


@login_required
def region_view(request, pk=None):
    obj = get_object_or_404(Region, pk=pk) if pk else None
    preview_data = None
    form = RegionForm(instance=obj)

    if request.method == "POST":
        # Prévisualisation du fichier CSV
        if "preview" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                df = pd.read_csv(fichier)
                preview_data = df.to_dict(orient="records")
            except Exception as e:
                messages.error(request, f"Erreur de lecture du fichier CSV : {e}")

        # Importation définitive du fichier CSV
        elif "import" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                df = pd.read_csv(fichier)
                for _, row in df.iterrows():
                    numero = row.get("numero")
                    nom = str(row.get("nom", "")).strip().upper()
                    if numero and nom:
                        Region.objects.get_or_create(numero=numero, nom=nom)
                messages.success(request, "Importation réussie !")
                return redirect("region")
            except Exception as e:
                messages.error(request, f"Erreur d'importation CSV : {e}")

        # Enregistrement manuel
        elif "save" in request.POST:
            form = RegionForm(request.POST, instance=obj)
            if form.is_valid():
                region = form.save(commit=False)
                region.nom = region.nom.strip().upper()  # Normalisation
                region.save()
                return redirect("region")
            else:
                messages.error(request, "Veuillez corriger les erreurs du formulaire.")

        # Suppression
        elif "delete" in request.POST and obj:
            obj.delete()
            return redirect("region")

    # Chargement standard
    regions = Region.objects.all()
    return render(
        request,
        "pages/region.html",
        {
            "form": form,
            "regions": regions,
            "obj": obj,
            "preview_data": preview_data,
        },
    )


@login_required
def region_detail(request, pk):
    region = get_object_or_404(Region, pk=pk)

    prefectures = region.prefectures.all().select_related("region")
    nombre_prefectures = prefectures.count()

    sous_prefectures = SousPrefecture.objects.filter(prefecture__region=region)
    nombre_sous_prefectures = sous_prefectures.count()

    # ✅ Récupération des communes
    communes = Commune.objects.filter(sous_prefecture__prefecture__region=region)
    nombre_communes = communes.count()

    try:
        prefecture_ids = list(map(int, request.GET.getlist("prefectures")))
    except (ValueError, TypeError):
        prefecture_ids = []

    try:
        secteur_ids = list(map(int, request.GET.getlist("secteurs")))
    except (ValueError, TypeError):
        secteur_ids = []

    try:
        communes_qs = communes  # Utilise les communes récupérées plus haut

        if prefecture_ids:
            communes_qs = communes_qs.filter(
                sous_prefecture__prefecture_id__in=prefecture_ids
            )

        centres = (
            CentreFormation.objects.filter(commune__in=communes_qs)
            .select_related("commune")
            .prefetch_related("domaine_activite_capacite__secteurs")
        )

        if secteur_ids:
            centres = centres.filter(secteurs__id__in=secteur_ids).distinct()

        nombre_centres = centres.count()

    except Exception as e:
        print(f"Erreur dans la requête: {e}")
        centres = CentreFormation.objects.none()
        nombre_centres = 0

    secteurs = Secteur.objects.filter(
        secteurs__centre__commune__sous_prefecture__prefecture__region=region
    ).distinct()

    data = []
    for pref in prefectures:
        sous_data = []
        for sous in pref.sousprefectures.all().select_related("prefecture"):
            commune_data = []
            for commune in sous.communes.all().select_related("sous_prefecture"):
                centres_commune = centres.filter(commune=commune)
                commune_data.append((commune, centres_commune))
            sous_data.append((sous, commune_data))
        data.append((pref, sous_data))

    return render(
        request,
        "pages/region_detail.html",
        {
            "region": region,
            "hierarchie": data,
            "nombre_centres": nombre_centres,
            "centres": centres,
            "prefecture_ids": prefecture_ids,
            "secteur_ids": secteur_ids,
            "secteurs": secteurs,
            "prefectures": prefectures,
            "nombre_prefectures": nombre_prefectures,
            "sous_prefectures": sous_prefectures,
            "nombre_sous_prefectures": nombre_sous_prefectures,
            "communes": communes,  # ✅ Ajout
            "nombre_communes": nombre_communes,  # ✅ Ajout
        },
    )


# View qui gere les prefectures
@login_required
def prefecture_view(request, pk=None):
    # Récupération de l'objet à modifier (s'il y a un pk dans l'URL)
    if pk:
        obj = get_object_or_404(Prefecture, pk=pk)
        form = PrefectureForm(request.POST or None, instance=obj)
    else:
        obj = None
        form = PrefectureForm(request.POST or None)

    if request.method == "POST":
        # Enregistrement ou modification
        if "save" in request.POST and form.is_valid():
            form.save()
            messages.success(request, "Préfecture enregistrée avec succès.")
            return redirect("prefecture")

        # Suppression
        elif "delete" in request.POST and obj:
            obj.delete()
            messages.success(request, "Préfecture supprimée avec succès.")
            return redirect("prefecture")

        # Prévisualisation CSV
        elif "preview" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                # Lecture du fichier CSV avec pandas
                df = pd.read_csv(fichier)

                # Enregistrement dans la session pour l'importation future
                request.session["df_prefecture"] = df.to_dict(orient="records")

                messages.info(
                    request,
                    "Prévisualisation chargée. Cliquez sur 'Importer' pour confirmer.",
                )
                return render(
                    request,
                    "pages/prefecture.html",
                    {
                        "form": form,
                        "prefectures": Prefecture.objects.all(),
                        "df_preview": df.to_dict(orient="records"),
                        "obj": obj,
                    },
                )
            except Exception as e:
                messages.error(
                    request, f"Erreur lors de la lecture du fichier CSV : {e}"
                )

        # Importation finale après prévisualisation
        elif "import" in request.POST:
            data = request.session.pop("df_prefecture", None)
            if data:
                for row in data:
                    try:
                        nom = row.get("nom")
                        region_nom = row.get("region")

                        if not nom or not region_nom:
                            messages.warning(
                                request, f"Ligne incomplète ignorée : {row}"
                            )
                            continue

                        # Recherche de la région par son nom (insensible à la casse)
                        try:
                            region_instance = Region.objects.get(
                                nom__iexact=region_nom.strip()
                            )
                        except Region.DoesNotExist:
                            messages.warning(
                                request,
                                f"Région '{region_nom}' non trouvée. Ligne ignorée.",
                            )
                            continue

                        # Création de la préfecture
                        Prefecture.objects.create(
                            nom=nom.strip(), region=region_instance
                        )

                    except Exception as e:
                        messages.error(
                            request, f"Erreur lors de l’importation d’une ligne : {e}"
                        )

                messages.success(request, "Importation terminée avec succès.")
            else:
                messages.error(request, "Aucune donnée disponible pour l'importation.")
            return redirect("prefecture")

    # Rendu de la page principale (GET ou retour après action)
    return render(
        request,
        "pages/prefecture.html",
        {
            "form": form,
            "prefectures": Prefecture.objects.all(),
            "obj": obj,
        },
    )


@login_required
def prefecture_detail(request, pk):
    prefecture = get_object_or_404(Prefecture, pk=pk)

    # ✅ Correction ici
    sous_prefectures = prefecture.sousprefectures.all().select_related("prefecture")
    nombre_sousprefecture = sous_prefectures.count()

    communes = Commune.objects.filter(sous_prefecture__prefecture=prefecture)
    nombre_communes = communes.count()

    try:
        sous_prefecture_ids = list(map(int, request.GET.getlist("sous_prefectures")))
    except (ValueError, TypeError):
        sous_prefecture_ids = []

    try:
        secteur_ids = list(map(int, request.GET.getlist("secteurs")))
    except (ValueError, TypeError):
        secteur_ids = []

    try:
        communes_qs = Commune.objects.filter(
            sous_prefecture__prefecture=prefecture
        ).select_related("sous_prefecture__prefecture")

        if sous_prefecture_ids:
            communes_qs = communes_qs.filter(sous_prefecture_id__in=sous_prefecture_ids)

        centres = (
            CentreFormation.objects.filter(commune__in=communes_qs)
            .select_related("commune")
            .prefetch_related("domaine_activite_capacite__secteurs")
        )

        if secteur_ids:
            centres = centres.filter(secteurs__id__in=secteur_ids).distinct()

        nombre_centres = centres.count()

    except Exception as e:
        print(f"Erreur dans la requête: {e}")
        centres = CentreFormation.objects.none()
        nombre_centres = 0

    secteurs = Secteur.objects.filter(
        secteurs__centre__commune__sous_prefecture__prefecture=prefecture
    ).distinct()

    data = []
    for sous in sous_prefectures:
        commune_data = []
        for commune in sous.communes.all().select_related("sous_prefecture"):
            centres_commune = centres.filter(commune=commune)
            commune_data.append((commune, centres_commune))
        data.append((sous, commune_data))

    return render(
        request,
        "pages/prefecture_detail.html",
        {
            "prefecture": prefecture,
            "hierarchie": data,
            "nombre_centres": nombre_centres,
            "centres": centres,
            "sous_prefecture_ids": sous_prefecture_ids,
            "secteur_ids": secteur_ids,
            "secteurs": secteurs,
            "sous_prefectures": sous_prefectures,
            "nombre_sousprefecture": nombre_sousprefecture,
            "communes": communes,
            "nombre_communes": nombre_communes,
        },
    )


# Vue qui gere les Sous Prefectures
@login_required
def sousprefecture_view(request, pk=None):
    if pk:
        obj = get_object_or_404(SousPrefecture, pk=pk)
        form = SousPrefectureForm(request.POST or None, instance=obj)
    else:
        obj = None
        form = SousPrefectureForm(request.POST or None)

    if request.method == "POST":
        # Enregistrement manuel
        if "save" in request.POST and form.is_valid():
            form.save()
            return redirect("sous_prefecture")

        # Suppression
        elif "delete" in request.POST and obj:
            obj.delete()
            return redirect("sous_prefecture")

        # Prévisualisation fichier CSV
        elif "preview" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                df = pd.read_csv(fichier)
                request.session["df_sousprefecture"] = df.to_dict(orient="records")
                messages.info(
                    request,
                    "Prévisualisation chargée. Cliquez sur 'Importer' pour confirmer.",
                )
                return render(
                    request,
                    "pages/sousprefecture.html",
                    {
                        "form": form,
                        "sousprefectures": SousPrefecture.objects.all(),
                        "df_preview": df.to_dict(orient="records"),
                        "obj": obj,
                    },
                )
            except Exception as e:
                messages.error(request, f"Erreur lecture fichier CSV : {e}")

        # Importation depuis la session
        elif "import" in request.POST:
            data = request.session.pop("df_sousprefecture", None)
            if data:
                for row in data:
                    try:
                        nom = row.get("nom")
                        prefecture_nom = row.get("prefecture") or row.get(
                            "prefecture_nom"
                        )
                        if not prefecture_nom:
                            messages.warning(
                                request,
                                f"Préfecture manquante pour la sous-préfecture '{nom}'. Ligne ignorée.",
                            )
                            continue
                        prefecture = Prefecture.objects.get(
                            nom__iexact=prefecture_nom.strip()
                        )
                        SousPrefecture.objects.create(
                            nom=nom.strip(), prefecture=prefecture
                        )
                    except Prefecture.DoesNotExist:
                        messages.warning(
                            request,
                            f"Préfecture '{prefecture_nom}' non trouvée pour la sous-préfecture '{nom}'. Ligne ignorée.",
                        )
                    except Exception as e:
                        messages.error(
                            request, f"Erreur import ligne pour '{nom}' : {e}"
                        )
                messages.success(request, "Importation terminée.")
            else:
                messages.error(request, "Aucune donnée à importer.")
            return redirect("sous_prefecture")

    return render(
        request,
        "pages/sousprefecture.html",
        {
            "form": form,
            "sousprefectures": SousPrefecture.objects.all(),
            "obj": obj,
        },
    )


@login_required
def sousprefecture_detail(request, pk):
    sous_prefecture = get_object_or_404(SousPrefecture, pk=pk)

    # ✅ Correction ici
    communes = sous_prefecture.communes.all().select_related("sous_prefecture")
    nombre_communes = communes.count()

    try:
        commune_ids = list(map(int, request.GET.getlist("communes")))
    except (ValueError, TypeError):
        commune_ids = []

    try:
        secteur_ids = list(map(int, request.GET.getlist("secteurs")))
    except (ValueError, TypeError):
        secteur_ids = []

    try:
        communes_qs = communes
        if commune_ids:
            communes_qs = communes_qs.filter(id__in=commune_ids)

        centres = (
            CentreFormation.objects.filter(commune__in=communes_qs)
            .select_related("commune")
            .prefetch_related("domaine_activite_capacite__secteurs")
        )

        if secteur_ids:
            centres = centres.filter(secteurs__id__in=secteur_ids).distinct()

        nombre_centres = centres.count()

    except Exception as e:
        print(f"Erreur dans la requête: {e}")
        centres = CentreFormation.objects.none()
        nombre_centres = 0

    secteurs = Secteur.objects.filter(
        secteurs__centre__commune__sous_prefecture=sous_prefecture
    ).distinct()

    data = []
    for commune in communes:
        centres_commune = centres.filter(commune=commune)
        data.append((commune, centres_commune))

    return render(
        request,
        "pages/sousprefecture_detail.html",
        {
            "sous_prefecture": sous_prefecture,
            "hierarchie": data,
            "nombre_centres": nombre_centres,
            "centres": centres,
            "commune_ids": commune_ids,
            "secteur_ids": secteur_ids,
            "secteurs": secteurs,
            "communes": communes,
            "nombre_communes": nombre_communes,
        },
    )


# Vue qui gere les Communes
@login_required
def commune_view(request, pk=None):
    if pk:
        obj = get_object_or_404(Commune, pk=pk)
        form = CommuneForm(request.POST or None, instance=obj)
    else:
        obj = None
        form = CommuneForm(request.POST or None)

    if request.method == "POST":
        # Enregistrement manuel
        if "save" in request.POST and form.is_valid():
            form.save()
            return redirect("commune")

        # Suppression
        elif "delete" in request.POST and obj:
            obj.delete()
            return redirect("commune")

        # Prévisualisation fichier CSV
        elif "preview" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                df = pd.read_csv(fichier)
                request.session["df_commune"] = df.to_dict(orient="records")
                messages.info(
                    request,
                    "Prévisualisation chargée. Cliquez sur 'Importer' pour confirmer.",
                )
                return render(
                    request,
                    "pages/commune.html",
                    {
                        "form": form,
                        "communes": Commune.objects.all(),
                        "df_preview": df.to_dict(orient="records"),
                        "obj": obj,
                    },
                )
            except Exception as e:
                messages.error(request, f"Erreur lecture fichier CSV : {e}")

        # Importation depuis la session
        elif "import" in request.POST:
            data = request.session.pop("df_commune", None)
            if data:
                for row in data:
                    try:
                        nom = row.get("nom")
                        sous_prefecture_nom = row.get("sous_prefecture") or row.get(
                            "sous_prefecture_nom"
                        )
                        if not sous_prefecture_nom:
                            messages.warning(
                                request,
                                f"Sous-préfecture manquante pour la commune '{nom}'. Ligne ignorée.",
                            )
                            continue
                        sous_prefecture = SousPrefecture.objects.get(
                            nom__iexact=sous_prefecture_nom.strip()
                        )
                        Commune.objects.create(
                            nom=nom.strip(), sous_prefecture=sous_prefecture
                        )
                    except SousPrefecture.DoesNotExist:
                        messages.warning(
                            request,
                            f"Sous-préfecture '{sous_prefecture_nom}' non trouvée pour la commune '{nom}'. Ligne ignorée.",
                        )
                    except Exception as e:
                        messages.error(
                            request, f"Erreur import ligne pour '{nom}' : {e}"
                        )
                messages.success(request, "Importation terminée.")
            else:
                messages.error(request, "Aucune donnée à importer.")
            return redirect("commune")

    return render(
        request,
        "pages/commune.html",
        {
            "form": form,
            "communes": Commune.objects.all(),
            "obj": obj,
        },
    )


@login_required
def commune_detail(request, pk):
    commune = get_object_or_404(Commune, pk=pk)

    secteur_ids = list(map(int, request.GET.getlist("secteurs")))
    centres = CentreFormation.objects.filter(commune=commune)
    if secteur_ids:
        centres = centres.filter(secteurs__id__in=secteur_ids).distinct()
    nombre_centres = centres.count()

    secteurs = Secteur.objects.all()

    # Préremplir le champ commune dans le formulaire CentreFormation
    initial_data = {"commune": commune}
    form_centre = CentreFormationForm(request.POST or None, initial=initial_data)
    form_docs = DocumentAdministratifForm(request.POST or None, request.FILES or None)
    form_ref = PersonneReferenceForm(request.POST or None)

    # Rendre le champ commune non modifiable dans le formulaire (readonly dans la vue)
    if request.method != "POST":
        if "commune" in form_centre.fields:
            form_centre.fields["commune"].widget.attrs["readonly"] = True
            form_centre.fields["commune"].widget.attrs["disabled"] = True

    # Traitement de la soumission
    if request.method == "POST":
        if form_centre.is_valid() and form_docs.is_valid() and form_ref.is_valid():
            centre = form_centre.save(commit=False)
            centre.commune = commune  # Associe automatiquement la commune
            centre.save()
            form_centre.save_m2m()  # Pour les champs ManyToMany comme secteurs

            doc = form_docs.save(commit=False)
            doc.centre = centre
            doc.save()

            ref = form_ref.save(commit=False)
            ref.centre = centre
            ref.save()

            messages.success(request, "Centre ajouté avec succès.")
            return redirect("commune_detail", pk=commune.pk)

    return render(
        request,
        "pages/commune_detail.html",
        {
            "commune": commune,
            "centres": centres,
            "nombre_centres": nombre_centres,
            "secteurs": secteurs,
            "secteur_ids": secteur_ids,
            "form_centre": form_centre,
            "form_docs": form_docs,
            "form_ref": form_ref,
        },
    )


# Vue qui gere les Secteurs
@login_required
def secteur_view(request, pk=None):
    if pk:
        obj = get_object_or_404(Secteur, pk=pk)
        form = SecteurForm(request.POST or None, instance=obj)
    else:
        obj = None
        form = SecteurForm(request.POST or None)

    if request.method == "POST":
        # Enregistrement manuel
        if "save" in request.POST and form.is_valid():
            form.save()
            return redirect("secteur")

        # Suppression
        elif "delete" in request.POST and obj:
            obj.delete()
            return redirect("secteur")

        # Prévisualisation fichier CSV
        elif "preview" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                import pandas as pd

                df = pd.read_csv(fichier)

                # Vérifie que la colonne 'nom' existe
                if "nom" not in df.columns:
                    messages.error(
                        request, "La colonne 'nom' est manquante dans le fichier CSV."
                    )
                else:
                    preview_data = df.to_dict(orient="records")
                    request.session["df_secteur"] = preview_data
                    messages.info(request, "Prévisualisation chargée.")
                    return render(
                        request,
                        "pages/secteur.html",
                        {
                            "form": form,
                            "secteurs": Secteur.objects.all(),
                            "df_preview": preview_data,
                            "obj": obj,
                        },
                    )
            except Exception as e:
                messages.error(request, f"Erreur lecture fichier CSV : {e}")

        # Importation depuis la session
        elif "import" in request.POST:
            data = request.session.pop("df_secteur", None)
            if data:
                for row in data:
                    try:
                        Secteur.objects.create(**row)
                    except Exception as e:
                        messages.error(request, f"Erreur import ligne : {e}")
                messages.success(request, "Importation réussie !")
            else:
                messages.error(request, "Aucune donnée à importer.")
            return redirect("secteur")

    return render(
        request,
        "pages/secteur.html",
        {
            "form": form,
            "secteurs": Secteur.objects.all(),
            "obj": obj,
        },
    )

@login_required
def publiccible_view(request, pk=None):
    if pk:
        obj = get_object_or_404(PublicCible, pk=pk)
        form = PublicCibleForm(request.POST or None, instance=obj)
    else:
        obj = None
        form = PublicCibleForm(request.POST or None)

    # Prévisualisation
    preview_data = None
    if request.method == "POST":
        # Enregistrement manuel
        if "save" in request.POST and form.is_valid():
            form.save()
            messages.success(request, "Public cible enregistré avec succès.")
            return redirect("publiccible")

        # Suppression
        elif "delete" in request.POST and obj:
            obj.delete()
            messages.success(request, "Public cible supprimé.")
            return redirect("publiccible")

        # Prévisualisation CSV
        elif "preview" in request.POST and "fichier" in request.FILES:
            fichier = request.FILES["fichier"]
            try:
                df = pd.read_csv(fichier, encoding="utf-8")
                preview_data = df.to_dict(orient="records")
                request.session["df_publiccible"] = preview_data
                messages.info(request, "Prévisualisation chargée.")
            except Exception as e:
                messages.error(request, f"Erreur lors de la lecture du fichier : {e}")

        # Importation depuis session
        elif "import" in request.POST:
            data = request.session.pop("df_publiccible", None)
            if data:
                success = 0
                errors = 0
                for row in data:
                    try:
                        PublicCible.objects.create(**row)
                        success += 1
                    except Exception as e:
                        errors += 1
                        messages.error(request, f"Erreur d'import : {e}")
                if success:
                    messages.success(
                        request, f"{success} public(s) cible importé(s) avec succès."
                    )
                if errors:
                    messages.warning(request, f"{errors} ligne(s) en erreur.")
            else:
                messages.error(request, "Aucune donnée disponible pour l'importation.")
            return redirect("publiccible")

    return render(
        request,
        "pages/publiccible.html",
        {
            "form": form,
            "publiccibles": PublicCible.objects.all(),
            "obj": obj,
            "preview_data": preview_data,
        },
    )

"""

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Exists, OuterRef
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from docx import Document
import pdfkit


@login_required
def centre_formation_view(request):
    print("🔍 Méthode requête :", request.method)
    print("📦 POST data :", request.POST)

    # Chargement des centres avec relations
    centres = CentreFormation.objects.prefetch_related(
        "secteurs", "domaineactivitecapacite_set"
    ).select_related("commune__sousprefecture__prefecture__region")

    # ✅ Annotation : un centre est "en règle" s’il a un fichier immatriculation_acfpe ET agrement_valide
    centres = centres.annotate(
        en_regle=Exists(
            DocumentAdministratif.objects.filter(
                centre=OuterRef('pk'),
                immatriculation_acfpe__isnull=False,
                agrement_valide__isnull=False
            )
        )
    )

    # Export PDF
    if request.GET.get("export") == "pdf":
        html = render_to_string("pages/centre_formation_pdf.html", {"centres": centres})
        pdf = pdfkit.from_string(html, False)
        response = HttpResponse(pdf, content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="centres_formation.pdf"'
        return response

    # Export Word
    if request.GET.get("export") == "word":
        doc = Document()
        doc.add_heading("Liste des Centres de Formation", 0)
        for centre in centres:
            doc.add_heading(centre.intitule, level=1)
            doc.add_paragraph(f"Commune : {centre.commune.nom}")
            secteurs_text = ", ".join(s.nom for s in centre.secteurs.all())
            doc.add_paragraph(f"Secteurs : {secteurs_text}")
            doc.add_paragraph(f"Adresse : {centre.adresse}")
            doc.add_paragraph(f"Téléphone : {centre.telephone}")
            doc.add_paragraph("")

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            'attachment; filename="centres_formation.docx"'
        )
        doc.save(response)
        return response

    # Formulaires
    form_centre = CentreFormationForm(request.POST or None)
    form_domaine = DomaineActiviteCapaciteForm(request.POST or None)
    form_formateur = FormateurForm(request.POST or None)
    form_docs = DocumentAdministratifForm(request.POST or None, request.FILES or None)
    form_ref = PersonneReferenceForm(request.POST or None)

    if request.method == "POST":
        if (
            form_centre.is_valid()
            and form_domaine.is_valid()
            and form_formateur.is_valid()
            and form_docs.is_valid()
            and form_ref.is_valid()
        ):
            centre = form_centre.save()

            domaine = form_domaine.save(commit=False)
            domaine.centre = centre
            domaine.save()
            form_domaine.save_m2m()

            formateur = form_formateur.save(commit=False)
            formateur.centre = centre
            formateur.save()

            doc = form_docs.save(commit=False)
            doc.centre = centre
            doc.save()

            ref = form_ref.save(commit=False)
            ref.centre = centre
            ref.save()

            messages.success(request, "Les informations ont été enregistrées avec succès !")
            return redirect("centre_formation")
        else:
            messages.error(request, "Erreur dans les formulaires. Veuillez corriger les champs.")
            print("Erreurs validation centre :", form_centre.errors)
            print("Erreurs validation domaine :", form_domaine.errors)
            print("Erreurs validation formateur :", form_formateur.errors)
            print("Erreurs validation docs :", form_docs.errors)
            print("Erreurs validation ref :", form_ref.errors)

    # 🔽 Préparation des données pour les filtres JS
    regions = Region.objects.all()
    prefectures = Prefecture.objects.all()
    sousprefectures = SousPrefecture.objects.all()
    communes = Commune.objects.all()
    secteurs = Secteur.objects.all()
    categories = CentreFormation.CATEGORIE_CHOICES  # ✅ labels des catégories

    return render(
        request,
        "pages/cf.html",
        {
            "form_centre": form_centre,
            "form_domaineActiviteCapacite": form_domaine,
            "form_formateur": form_formateur,
            "form_docs": form_docs,
            "form_ref": form_ref,
            "centres": centres,
            "document": DocumentAdministratif.objects.all(),
            "regions": regions,
            "prefectures": prefectures,
            "sousprefectures": sousprefectures,
            "communes": communes,
            "secteurs": secteurs,
            "categories": categories,  # ✅ transmis avec labels
        },
    )





"""


@login_required
def centre_formation_view(request):
    print("🔍 Méthode requête :", request.method)
    print("📦 POST data :", request.POST)
    is_gestionnaire = request.user.groups.filter(name='gestionnaire').exists()
    centres = list(CentreFormation.objects.all())
    categories = CentreFormation.CATEGORIE_CHOICES

    # ✅ Ajouter l'attribut dynamique `en_regle` à chaque centre
    for centre in centres:
        centre.en_regle = DocumentAdministratif.objects.filter(
            centre=centre,
            immatriculation_acfpe__isnull=False,
            agrement_valide__isnull=False
        ).exclude(
            immatriculation_acfpe="",
            agrement_valide=""
        ).exists()
        print(f"[DEBUG] {centre.intitule} - En règle : {centre.en_regle}")
    # Export PDF
    if request.GET.get("export") == "pdf":
        html = render_to_string("pages/centre_formation_pdf.html", {"centres": centres})
        pdf = pdfkit.from_string(html, False)
        response = HttpResponse(pdf, content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="centres_formation.pdf"'
        return response

    # Export Word
    if request.GET.get("export") == "word":
        doc = Document()
        doc.add_heading("Liste des Centres de Formation", 0)
        for centre in centres:
            doc.add_heading(centre.intitule, level=1)
            doc.add_paragraph(f"Commune : {centre.commune.nom}")
            secteurs_text = ", ".join(s.nom for s in centre.secteurs.all())
            doc.add_paragraph(f"Secteurs : {secteurs_text}")
            doc.add_paragraph(f"Adresse : {centre.adresse}")
            doc.add_paragraph(f"Téléphone : {centre.telephone}")
            doc.add_paragraph("")
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            'attachment; filename="centres_formation.docx"'
        )
        doc.save(response)
        return response

    # Formulaires
    form_centre = CentreFormationForm(request.POST or None)
    form_domaine = DomaineActiviteCapaciteForm(request.POST or None)
    form_formateur = FormateurForm(request.POST or None)
    form_docs = DocumentAdministratifForm(request.POST or None, request.FILES or None)
    form_ref = PersonneReferenceForm(request.POST or None)

    if request.method == "POST":
        if (
            form_centre.is_valid()
            and form_domaine.is_valid()
            and form_formateur.is_valid()
            and form_docs.is_valid()
            and form_ref.is_valid()
        ):
            centre = form_centre.save()

            domaine = form_domaine.save(commit=False)
            domaine.centre = centre
            domaine.save()
            form_domaine.save_m2m()

            formateur = form_formateur.save(commit=False)
            formateur.centre = centre
            formateur.save()

            doc = form_docs.save(commit=False)
            doc.centre = centre
            doc.save()

            ref = form_ref.save(commit=False)
            ref.centre = centre
            ref.save()

            messages.success(
                request, "Les informations ont été enregistrées avec succès !"
            )
            return redirect("centre_formation")
        else:
            messages.error(
                request, "Erreur dans les formulaires. Veuillez corriger les champs."
            )
            print("Erreurs validation centre :", form_centre.errors)
            print("Erreurs validation domaine :", form_domaine.errors)
            print("Erreurs validation formateur :", form_formateur.errors)
            print("Erreurs validation docs :", form_docs.errors)
            print("Erreurs validation ref :", form_ref.errors)

    return render(
        request,
        "pages/cf.html",
        {
            "form_centre": form_centre,
            "form_domaineActiviteCapacite": form_domaine,
            "form_formateur": form_formateur,
            "form_docs": form_docs,
            "form_ref": form_ref,
            "centres": centres,
            "document": DocumentAdministratif.objects.all(),
            "regions": Region.objects.all(),
            "prefectures": Prefecture.objects.all(),
            "sousprefectures": SousPrefecture.objects.all(),
            "communes": Commune.objects.all(),
            "secteurs": Secteur.objects.all(),
            "categories": categories,
            "is_gestionnaire": is_gestionnaire,
    "can_add_centre": True,  # Les gestionnaires peuvent ajouter
        },
    )


@login_required
def centre_detail(request, pk):

    is_gestionnaire = request.user.groups.filter(name='gestionnaire').exists()
    centre = get_object_or_404(
        CentreFormation.objects.select_related(
            "commune__sous_prefecture__prefecture__region",
            "domaine_activite_capacite",
            "document_administratif",
            "personne_reference",
        ).prefetch_related(
            "domaine_activite_capacite__secteurs",
            "domaine_activite_capacite__public_cibles",
        ),
        pk=pk,
    )

    try:
        doc = centre.document_administratif
    except DocumentAdministratif.DoesNotExist:
        doc = DocumentAdministratif(centre=centre)

    try:
        ref = centre.personne_reference
    except PersonneReference.DoesNotExist:
        ref = PersonneReference(centre=centre)

    try:
        formateur = centre.formateur
    except Formateur.DoesNotExist:
        formateur = Formateur(centre=centre)

    # 🔹 Export PDF
    if request.GET.get("export") == "pdf":
        template_path = "pages/centre_pdf.html"
        context = {
            "centre": centre,
            "commune": centre.commune,
            "sous_prefecture": centre.commune.sous_prefecture,
            "prefecture": centre.commune.sous_prefecture.prefecture,
            "region": centre.commune.sous_prefecture.prefecture.region,
            "secteurs": (
                centre.domaine_activite_capacite.secteurs.all()
                if hasattr(centre, "domaine_activite_capacite")
                else []
            ),
            "publics_cibles": (
                centre.domaine_activite_capacite.public_cibles.all()
                if hasattr(centre, "domaine_activite_capacite")
                else []
            ),
            "documents": doc,
            "personne_reference": ref,
            "formateur": formateur,
            "niveaux_formateur": (
                formateur.niveaux_formateur
                if formateur and formateur.niveaux_formateur
                else []
            ),
            "experience_formateur": (
                formateur.experience_formateur
                if formateur and formateur.experience_formateur
                else []
            ),
            "nombre_formateur_permanant": (
                formateur.nombre_formateur_permanant if formateur else 0
            ),
            "nombre_formateur_nonpermanant": (
                formateur.nombre_formateur_nonpermanant if formateur else 0
            ),
        }
        html = render_to_string(template_path, context)
        pdf = pdfkit.from_string(html, False)
        response = HttpResponse(pdf, content_type="application/pdf")
        response["Content-Disposition"] = (
            f'attachment; filename="centre_{centre.intitule}.pdf"'
        )
        return response  # ✅ maintenant dans le bloc if

    # 🔹 Traitement des formulaires
    if request.method == "POST":
        if is_gestionnaire:
            messages.error(request, "Vous n'avez pas les permissions pour modifier un centre.")
            return redirect("centre_detail", pk=centre.pk)
        if "delete" in request.POST:
            centre.delete()
            messages.success(request, "Centre supprimé avec succès.")
            return redirect("centre_formation")

        form_centre = CentreFormationForm(request.POST, instance=centre)
        form_domaineActiviteCapacite = DomaineActiviteCapaciteForm(
            request.POST,
            instance=(
                centre.domaine_activite_capacite
                if hasattr(centre, "domaine_activite_capacite")
                else None
            ),
        )
        form_formateur = FormateurForm(request.POST, instance=formateur)
        form_docs = DocumentAdministratifForm(request.POST, request.FILES, instance=doc)
        form_ref = PersonneReferenceForm(request.POST, instance=ref)

        if all(
            [
                form_centre.is_valid(),
                form_domaineActiviteCapacite.is_valid(),
                form_formateur.is_valid(),
                form_docs.is_valid(),
                form_ref.is_valid(),
            ]
        ):
            form_centre.save()
            form_domaineActiviteCapacite.instance.centre = centre
            form_domaineActiviteCapacite.save()
            form_formateur.instance.centre = centre
            form_formateur.save()
            form_docs.instance.centre = centre
            form_docs.save()
            form_ref.instance.centre = centre
            form_ref.save()
            messages.success(request, "Modifications enregistrées avec succès.")
            return redirect("centre_detail", pk=centre.pk)
    else:
        form_centre = CentreFormationForm(instance=centre)
        form_domaineActiviteCapacite = DomaineActiviteCapaciteForm(
            instance=(
                centre.domaine_activite_capacite
                if hasattr(centre, "domaine_activite_capacite")
                else None
            )
        )
        form_formateur = FormateurForm(instance=formateur)
        form_docs = DocumentAdministratifForm(instance=doc)
        form_ref = PersonneReferenceForm(instance=ref)

    return render(
        request,
        "pages/cf_detail.html",
        {
            "centre": centre,
            "form_centre": form_centre,
            "form_domaineActiviteCapacite": form_domaineActiviteCapacite,
            "form_formateur": form_formateur,
            "form_docs": form_docs,
            "form_ref": form_ref,
            "commune": centre.commune,
            "sous_prefecture": centre.commune.sous_prefecture,
            "prefecture": centre.commune.sous_prefecture.prefecture,
            "region": centre.commune.sous_prefecture.prefecture.region,
            "secteurs": (
                centre.domaine_activite_capacite.secteurs.all()
                if hasattr(centre, "domaine_activite_capacite")
                else []
            ),
            "publics_cibles": (
                centre.domaine_activite_capacite.public_cibles.all()
                if hasattr(centre, "domaine_activite_capacite")
                else []
            ),
            "niveaux_formateur": (
                formateur.niveaux_formateur
                if formateur and formateur.niveaux_formateur
                else []
            ),
            "experience_formateur": (
                formateur.experience_formateur
                if formateur and formateur.experience_formateur
                else []
            ),
            "documents": doc,
            "personne_reference": ref,
             "is_gestionnaire": is_gestionnaire,  # AJOUTÉ AU CONTEXTE
            "can_edit": not is_gestionnaire,     # AJOUTÉ AU CONTEXTE
            "can_delete": not is_gestionnaire,   # AJO
        },
    )

@login_required
def index(request):
    # 📊 Données géographiques
    region_data = list(
        Region.objects.annotate(
            total=Count("prefectures__sousprefectures__communes__centres")
        ).values("nom", "total")
    )
    pref_data = list(
        Prefecture.objects.annotate(
            total=Count("sousprefectures__communes__centres")
        ).values("nom", "total")
    )
    souspref_data = list(
        SousPrefecture.objects.annotate(
            total=Count("communes__centres")
        ).values("nom", "total")
    )
    commune_data = list(
        Commune.objects.annotate(
            total=Count("centres")
        ).values("nom", "total")
    )

    # 📋 Statistiques globales
    total_regions = Region.objects.count()
    total_prefectures = Prefecture.objects.count()
    total_sousprefectures = SousPrefecture.objects.count()
    total_communes = Commune.objects.count()
    total_secteurs = Secteur.objects.count()
    total_centre = CentreFormation.objects.count()

    # 🔄 Centres unisectoriels et multisectoriels
    total_centres_uni = DomaineActiviteCapacite.objects.annotate(
        nb_secteurs=Count("secteurs")
    ).filter(nb_secteurs=1).count()

    total_centres_multi = DomaineActiviteCapacite.objects.annotate(
        nb_secteurs=Count("secteurs")
    ).filter(nb_secteurs__gt=1).count()

    # ✅ Centres en règle
    centres_en_regle = CentreFormation.objects.filter(
        document_administratif__isnull=False,
        document_administratif__immatriculation_acfpe__isnull=False,
    ).exclude(
        document_administratif__immatriculation_acfpe=""
    ).filter(
        document_administratif__agrement_valide__isnull=False,
    ).exclude(
        document_administratif__agrement_valide=""
    )

    total_centre_en_regle = centres_en_regle.count()
    total_centre_non_regle = total_centre - total_centre_en_regle

    # 📊 Centres par catégorie
    categorie_data = list(
        CentreFormation.objects.values("categorie")
        .annotate(total=Count("id"))
    )

    # Mapping des catégories (1e, 2e, 3e) vers labels lisibles
    categorie_mapping = dict(CentreFormation.CATEGORIE_CHOICES)
    for c in categorie_data:
        c["categorie"] = categorie_mapping.get(c["categorie"], "Non spécifiée")

    # 🔄 Contexte à passer au template
    context = {
        # 📊 Données pour les graphiques géographiques
        "region_data": json.dumps(region_data),
        "prefecture_data": json.dumps(pref_data),
        "sousprefecture_data": json.dumps(souspref_data),
        "commune_data": json.dumps(commune_data),

        # 🧮 Statistiques générales
        "total_regions": total_regions,
        "total_prefectures": total_prefectures,
        "total_sousprefectures": total_sousprefectures,
        "total_communes": total_communes,
        "total_secteurs": total_secteurs,
        "total_centres_uni": total_centres_uni,
        "total_centres_multi": total_centres_multi,
        "total_centre": total_centre,
        "total_centre_en_regle": total_centre_en_regle,

        # 📊 Données pour graphiques par catégorie et régularité
        "categorie_data": json.dumps(categorie_data),
        "centre_regle_data": json.dumps({
            "en_regle": total_centre_en_regle,
            "non_regle": total_centre_non_regle,
        }),
    }

    return render(request, "pages/index.html", context)



# def calendar(request):
#    return render(request,"pages/index.html")
MESSAGE_TAGS = {
    messages.DEBUG: 'debug',
    messages.INFO: 'info',
    messages.SUCCESS: 'success',
    messages.WARNING: 'warning',
    messages.ERROR: 'error',
}

def login_view(request):
    if request.user.is_authenticated:
        return redirect('index')

    form = AuthenticationForm(request, data=request.POST or None)

    if request.method == 'POST':
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)

            if user is not None:
                login(request, user)
                messages.success(request, f"Connexion Reussie\n Bienvenue {user.username} !")
                return redirect('index')
            else:
                # Ce bloc est rarement atteint si form.is_valid() est True, mais on garde une sécurité
                messages.error(request, "Nom d'utilisateur ou mot de passe incorrect.")
        else:
            # Affichage de message global pour erreur de connexion
            messages.error(request, "Nom d'utilisateur ou mot de passe incorrect\n Veullez ressayer.")

    return render(request, 'pages/login-v1.html', {'form': form})



def logout_view(request):
    logout(request)
    messages.info(request, "Déconnexion réussie.")
    return redirect('login')
